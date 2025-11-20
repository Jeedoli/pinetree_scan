# RAG 시스템 핵심 모듈
import os
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
import re
from dataclasses import dataclass
from dotenv import load_dotenv

# 환경 변수 로드
load_dotenv()

# AI 모델 관련 import
try:
    from transformers import AutoTokenizer, AutoModelForCausalLM
    import torch
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False
    print("⚠️ transformers 라이브러리 없음 - LLM 기능 비활성화")

# LangChain + OpenAI 관련 import
try:
    from langchain_openai import ChatOpenAI
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_core.output_parsers import StrOutputParser
    LANGCHAIN_AVAILABLE = True
    print("✅ LangChain + OpenAI 라이브러리 로드 성공")
except ImportError as e:
    LANGCHAIN_AVAILABLE = False
    print(f"⚠️ LangChain 라이브러리 없음: {e}")

# Semantic Search 관련 import
try:
    from sentence_transformers import SentenceTransformer
    import faiss
    import numpy as np
    SEMANTIC_SEARCH_AVAILABLE = True
except ImportError:
    SEMANTIC_SEARCH_AVAILABLE = False
    print("⚠️ semantic search 라이브러리 없음 (sentence-transformers, faiss-cpu)")

# 임베딩 캐시용
import pickle

@dataclass
class Document:
    """문서 객체"""
    content: str
    metadata: Dict[str, Any]
    filename: str = ""
    
class SimpleRAG:
    """Semantic RAG 시스템 - HuggingFace + FAISS 기반"""
    
    def __init__(self, knowledge_base_path: str):
        # MPS 완전 비활성화 (MacOS M1/M2/M3)
        import os
        os.environ['PYTORCH_ENABLE_MPS_FALLBACK'] = '0'
        os.environ['PYTORCH_MPS_HIGH_WATERMARK_RATIO'] = '0.0'
        
        self.knowledge_base_path = Path(knowledge_base_path)
        self.documents: List[Document] = []
        self.tokenizer = None
        self.model = None
        
        # Semantic Search 컴포넌트
        self.embedding_model = None
        self.faiss_index = None
        self.document_embeddings = None
        self.embedding_cache_path = self.knowledge_base_path.parent / "embeddings_cache.pkl"
        
        # OpenAI LLM
        self.openai_llm = None
        self.openai_chain = None
        
        self.load_knowledge_base()
        self.load_embedding_model()
        self.build_semantic_index()
        self.setup_openai_llm()
    
    def load_embedding_model(self):
        """임베딩 모델 로드"""
        if not SEMANTIC_SEARCH_AVAILABLE:
            print("⚠️ semantic search 라이브러리가 없어 키워드 검색을 사용합니다")
            return
            
        try:
            print("🤗 Sentence Transformer 모델 로딩...")
            # 한국어 + 영어 지원하는 multilingual 모델 사용
            self.embedding_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
            print("✅ 임베딩 모델 로드 완료")
        except Exception as e:
            print(f"⚠️ 임베딩 모델 로드 실패: {e}")
            self.embedding_model = None
    
    def build_semantic_index(self):
        """문서 임베딩 생성 및 FAISS 인덱스 구축"""
        if not self.embedding_model or not self.documents:
            print("⚠️ 임베딩 모델이 없거나 문서가 없어 인덱스를 구축할 수 없습니다")
            return
            
        print("🔍 문서 임베딩 및 FAISS 인덱스 구축 중...")
        
        # 캐시된 임베딩이 있는지 확인
        if self._load_cached_embeddings():
            print("✅ 캐시된 임베딩 사용")
            return
            
        try:
            # 각 문서의 텍스트를 청크로 나누기
            document_chunks = []
            chunk_to_doc_map = []
            
            for doc_idx, doc in enumerate(self.documents):
                # 문서를 512자 단위로 청크 분할 (오버랩 100자)
                chunks = self._split_document_into_chunks(doc.content, chunk_size=512, overlap=100)
                for chunk in chunks:
                    if len(chunk.strip()) > 50:  # 너무 짧은 청크 제외
                        document_chunks.append(chunk)
                        chunk_to_doc_map.append(doc_idx)
            
            if not document_chunks:
                print("⚠️ 유효한 문서 청크가 없습니다")
                return
                
            print(f"📄 총 {len(document_chunks)}개 청크 생성")
            
            # 임베딩 생성
            print("🔄 임베딩 생성 중...")
            embeddings = self.embedding_model.encode(
                document_chunks,
                show_progress_bar=True,
                batch_size=16
            )
            
            # FAISS 인덱스 구축
            dimension = embeddings.shape[1]
            self.faiss_index = faiss.IndexFlatIP(dimension)  # Inner Product (코사인 유사도)
            
            # 정규화 (코사인 유사도를 위해)
            faiss.normalize_L2(embeddings)
            self.faiss_index.add(embeddings.astype(np.float32))
            
            # 메타데이터 저장
            self.document_embeddings = {
                'chunks': document_chunks,
                'chunk_to_doc_map': chunk_to_doc_map,
                'embeddings': embeddings
            }
            
            # 캐시 저장
            self._save_embeddings_cache()
            
            print(f"✅ FAISS 인덱스 구축 완료 ({len(document_chunks)}개 청크, {dimension}차원)")
            
        except Exception as e:
            print(f"⚠️ 인덱스 구축 실패: {e}")
            self.faiss_index = None
            self.document_embeddings = None
    
    def _split_document_into_chunks(self, text: str, chunk_size: int = 512, overlap: int = 100) -> List[str]:
        """문서를 오버랩이 있는 청크로 분할"""
        if len(text) <= chunk_size:
            return [text]
            
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # 단어 경계에서 자르기
            chunk = text[start:end]
            last_space = chunk.rfind(' ')
            if last_space != -1 and last_space > chunk_size * 0.7:
                chunk = chunk[:last_space]
                end = start + last_space
            
            chunks.append(chunk)
            start = end - overlap
            
        return chunks
    
    def _load_cached_embeddings(self) -> bool:
        """캐시된 임베딩 로드"""
        try:
            if self.embedding_cache_path.exists():
                with open(self.embedding_cache_path, 'rb') as f:
                    cache_data = pickle.load(f)
                
                # 문서 개수가 같은지 확인
                if cache_data.get('doc_count') == len(self.documents):
                    self.document_embeddings = cache_data['embeddings']
                    
                    # FAISS 인덱스 복원
                    embeddings = cache_data['embeddings']['embeddings']
                    dimension = embeddings.shape[1]
                    self.faiss_index = faiss.IndexFlatIP(dimension)
                    self.faiss_index.add(embeddings.astype(np.float32))
                    
                    return True
            return False
        except Exception as e:
            print(f"⚠️ 임베딩 캐시 로드 실패: {e}")
            return False
    
    def _save_embeddings_cache(self):
        """임베딩 캐시 저장"""
        try:
            cache_data = {
                'doc_count': len(self.documents),
                'embeddings': self.document_embeddings
            }
            with open(self.embedding_cache_path, 'wb') as f:
                pickle.dump(cache_data, f)
            print(f"💾 임베딩 캐시 저장: {self.embedding_cache_path}")
        except Exception as e:
            print(f"⚠️ 임베딩 캐시 저장 실패: {e}")

    def load_knowledge_base(self):
        """지식 베이스 로드"""
        print("📚 지식 베이스 로딩 시작...")
        
        self.documents = []
        if not self.knowledge_base_path.exists():
            print("⚠️ 지식 베이스 디렉토리가 없습니다.")
            return
        
        # 지원하는 파일 형식들
        supported_extensions = ["*.md", "*.txt", "*.json", "*.csv"]
        
        # PDF 지원 (PyPDF2가 설치된 경우)
        try:
            import PyPDF2
            supported_extensions.append("*.pdf")
        except ImportError:
            print("💡 PyPDF2가 설치되지 않아 PDF 파일은 지원하지 않습니다.")
            print("   PDF 지원을 원한다면: poetry add PyPDF2")
        
        # 워드 문서 지원 (python-docx가 설치된 경우)  
        try:
            import docx
            supported_extensions.append("*.docx")
        except ImportError:
            print("💡 python-docx가 설치되지 않아 DOCX 파일은 지원하지 않습니다.")
            print("   DOCX 지원을 원한다면: poetry add python-docx")
        
        for pattern in supported_extensions:
            for file_path in self.knowledge_base_path.rglob(pattern):
                try:
                    content = self._load_file_content(file_path)
                    if content:
                        doc = Document(
                            content=content,
                            metadata={
                                'source_file': file_path.name,
                                'category': file_path.parent.name,
                                'file_path': str(file_path),
                                'file_type': file_path.suffix[1:]  # 확장자 저장
                            },
                            filename=file_path.name
                        )
                        self.documents.append(doc)
                        print(f"✅ 로드 완료: {file_path.name} ({file_path.suffix})")
                        
                except Exception as e:
                    print(f"⚠️ 파일 로드 실패: {file_path.name} - {e}")
        
        print(f"📖 총 {len(self.documents)}개 문서 로드 완료")
    
    def _load_file_content(self, file_path: Path) -> str:
        """파일 형식별 내용 로드"""
        try:
            if file_path.suffix.lower() == '.md':
                # 마크다운 파일
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif file_path.suffix.lower() == '.txt':
                # 텍스트 파일
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif file_path.suffix.lower() == '.json':
                # JSON 파일 - 구조화된 데이터를 텍스트로 변환
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return self._json_to_text(data, file_path.name)
            
            elif file_path.suffix.lower() == '.csv':
                # CSV 파일 - 표 형태 데이터를 텍스트로 변환
                try:
                    import pandas as pd
                    df = pd.read_csv(file_path)
                    return f"# {file_path.name}\n\n{df.to_string()}"
                except ImportError:
                    # pandas 없으면 기본 CSV 읽기
                    import csv
                    content = [f"# {file_path.name}\n"]
                    with open(file_path, 'r', encoding='utf-8') as f:
                        reader = csv.reader(f)
                        for row in reader:
                            content.append(" | ".join(row))
                    return "\n".join(content)
            
            elif file_path.suffix.lower() == '.pdf':
                # PDF 파일
                try:
                    import PyPDF2
                    content = [f"# {file_path.name}\n"]
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            content.append(page.extract_text())
                    return "\n".join(content)
                except ImportError:
                    print(f"⚠️ PyPDF2가 없어 PDF 파일을 읽을 수 없습니다: {file_path.name}")
                    return ""
            
            elif file_path.suffix.lower() == '.docx':
                # Word 문서
                try:
                    import docx
                    doc = docx.Document(file_path)
                    content = [f"# {file_path.name}\n"]
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            content.append(paragraph.text)
                    return "\n".join(content)
                except ImportError:
                    print(f"⚠️ python-docx가 없어 DOCX 파일을 읽을 수 없습니다: {file_path.name}")
                    return ""
            
            else:
                print(f"⚠️ 지원하지 않는 파일 형식: {file_path.suffix}")
                return ""
                
        except Exception as e:
            print(f"❌ 파일 읽기 오류 ({file_path.name}): {e}")
            return ""
    
    def _json_to_text(self, data: dict, filename: str) -> str:
        """JSON 데이터를 검색 가능한 텍스트로 변환"""
        text_parts = [f"# {filename}\n"]
        
        def extract_text(obj, prefix=""):
            if isinstance(obj, dict):
                for key, value in obj.items():
                    text_parts.append(f"{prefix}{key}: {value}")
                    if isinstance(value, (dict, list)):
                        extract_text(value, f"{prefix}  ")
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    text_parts.append(f"{prefix}[{i}]: {item}")
                    if isinstance(item, (dict, list)):
                        extract_text(item, f"{prefix}  ")
        
        extract_text(data)
        return "\n".join(text_parts)
    
    def load_model(self):
        """한국어 특화 LLM 로드 - 하드코딩 없이 자유롭게 대화"""
        if not TRANSFORMERS_AVAILABLE:
            print("⚠️ transformers 라이브러리가 없어 기본 모드로 실행됩니다.")
            return
            
        # 한국어 생성 가능 모델들
        model_candidates = [
            # 한국어 특화
            ("skt/kogpt2-base-v2", "SKT KoGPT-2 한국어 모델"),
            ("skt/ko-gpt-trinity-1.2B-v0.5", "KoGPT Trinity"),
            # 다국어 지원
            ("facebook/opt-125m", "OPT 125M"),
            ("EleutherAI/gpt-neo-125M", "GPT-Neo 125M"),
            # Fallback
            ("gpt2", "GPT-2"),
        ]
        
        for model_name, description in model_candidates:
            try:
                print(f"🧠 {description} 로딩 중... ({model_name})")
                
                self.tokenizer = AutoTokenizer.from_pretrained(
                    model_name,
                    trust_remote_code=True
                )
                # MacOS MPS 문제 회피 - CPU만 사용
                import os
                os.environ['PYTORCH_ENABLE_MPS_FALLBACK'] = '1'
                
                self.model = AutoModelForCausalLM.from_pretrained(
                    model_name,
                    torch_dtype=torch.float32,
                    trust_remote_code=True,
                    low_cpu_mem_usage=True
                )
                
                # CPU 강제
                self.model = self.model.to('cpu')
                self.model.eval()
                print("📍 모델을 CPU에서 실행합니다")
                
                # 패딩 토큰 설정
                if self.tokenizer.pad_token is None:
                    self.tokenizer.pad_token = self.tokenizer.eos_token
                    if self.model.config.pad_token_id is None:
                        self.model.config.pad_token_id = self.tokenizer.eos_token_id
                    
                print(f"✅ {description} 로드 성공! 자유로운 한국어 대화 가능")
                return
                
            except Exception as e:
                print(f"⚠️ {model_name} 로드 실패: {e}")
                continue
        
        print("❌ 모든 LLM 로드 실패. 지식베이스만 사용합니다.")
        self.tokenizer = None
        self.model = None
    
    def setup_openai_llm(self):
        """OpenAI GPT-3.5-turbo 설정"""
        if not LANGCHAIN_AVAILABLE:
            print("⚠️ LangChain이 설치되지 않아 OpenAI를 사용할 수 없습니다")
            return
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key or api_key == "your_openai_api_key_here":
            print("⚠️ OPENAI_API_KEY가 .env에 설정되지 않았습니다")
            print("💡 .env 파일을 생성하고 OPENAI_API_KEY를 설정해주세요")
            return
        
        try:
            model_name = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")
            
            # OpenAI LLM 초기화
            self.openai_llm = ChatOpenAI(
                model=model_name,
                temperature=0.7,  # 창의적이고 자연스러운 답변
                api_key=api_key
            )
            
            # RAG 프롬프트 템플릿
            prompt = ChatPromptTemplate.from_messages([
                ("system", """당신은 소나무재선충병 탐지 시스템 및 AI/머신러닝 전문가입니다.
주어진 정보와 전문 지식을 활용하여 정확하고 이해하기 쉽게 한국어로 답변해주세요.

답변 원칙:
1. Context에 직접적인 정보가 있으면 우선 활용
2. Context에 정보가 부족하면 전문 지식으로 보완하여 답변
3. 완전히 모르는 내용만 "정확한 정보를 찾지 못했습니다"
4. 전문 용어는 쉽게 설명
5. 실용적이고 구체적으로 2-4문장으로 답변
6. mAP, 신뢰도, 탐지 성능 등 기술적 질문에는 일반적인 산업 기준 제시"""),
                ("user", """참고 자료: {context}

질문: {question}

답변:""")
            ])
            
            # Chain 구성
            self.openai_chain = prompt | self.openai_llm | StrOutputParser()
            
            print(f"✅ OpenAI {model_name} 모델 설정 완료")
            
        except Exception as e:
            print(f"⚠️ OpenAI 설정 실패: {e}")
            self.openai_llm = None
            self.openai_chain = None
    
    def simple_search(self, query: str, top_k: int = 3) -> List[Document]:
        """Semantic 검색 (FAISS) 또는 키워드 기반 검색"""
        if not self.documents:
            print("📚 검색할 문서가 없습니다.")
            return []
        
        # Semantic Search 우선 시도
        if self.faiss_index is not None and self.embedding_model is not None:
            return self._semantic_search(query, top_k)
        else:
            print("🔄 Semantic search 불가, 키워드 검색 사용")
            return self._keyword_search(query, top_k)
    
    def _semantic_search(self, query: str, top_k: int = 3) -> List[Document]:
        """FAISS 기반 Semantic 검색"""
        try:
            print(f"🤗 Semantic 검색: '{query}' (총 {len(self.documents)}개 문서)")
            
            # 쿼리 임베딩 생성
            query_embedding = self.embedding_model.encode([query])
            faiss.normalize_L2(query_embedding)
            
            # FAISS 검색
            scores, indices = self.faiss_index.search(
                query_embedding.astype(np.float32), 
                k=min(top_k * 3, len(self.document_embeddings['chunks']))  # 더 많이 찾아서 다양성 확보
            )
            
            # 청크를 문서별로 그룹핑
            doc_scores = {}
            for score, idx in zip(scores[0], indices[0]):
                if idx == -1:  # 검색 실패
                    continue
                    
                doc_idx = self.document_embeddings['chunk_to_doc_map'][idx]
                chunk_text = self.document_embeddings['chunks'][idx]
                
                if doc_idx not in doc_scores:
                    doc_scores[doc_idx] = {
                        'max_score': score,
                        'avg_score': score,
                        'count': 1,
                        'best_chunk': chunk_text[:200] + '...'
                    }
                else:
                    doc_scores[doc_idx]['max_score'] = max(doc_scores[doc_idx]['max_score'], score)
                    doc_scores[doc_idx]['avg_score'] = (doc_scores[doc_idx]['avg_score'] * doc_scores[doc_idx]['count'] + score) / (doc_scores[doc_idx]['count'] + 1)
                    doc_scores[doc_idx]['count'] += 1
                    
                    # 더 높은 점수 청크로 대표 청크 업데이트
                    if score > doc_scores[doc_idx]['max_score'] * 0.95:
                        doc_scores[doc_idx]['best_chunk'] = chunk_text[:200] + '...'
            
            # 문서별 점수로 정렬 - 관련성을 더 중요하게 고려
            # 키워드 매칭 보너스 추가
            query_lower = query.lower()
            key_terms = ['재선충', '소나무재선충', 'pine wilt', '탐지', 'detection', '확인', '방법', '증상']
            
            for doc_idx, score_info in doc_scores.items():
                doc = self.documents[doc_idx]
                content_lower = doc.content.lower()
                
                # 키워드 매칭 보너스
                keyword_bonus = 0
                for term in key_terms:
                    if term in query_lower and term in content_lower:
                        keyword_bonus += 0.1
                
                # 파일명 관련성 보너스
                filename = doc.metadata.get('source_file', '').lower()
                if any(term in filename for term in ['재선충', '소나무', 'pine', '탐지', 'detection']):
                    keyword_bonus += 0.15
                
                # 설정 파일 페널티
                if 'setting' in filename or '설정' in filename or 'config' in filename:
                    keyword_bonus -= 0.3
                
                score_info['final_score'] = score_info['max_score'] + score_info['avg_score'] + keyword_bonus
            
            # 최종 점수로 정렬 (final_score 우선)
            sorted_docs = sorted(doc_scores.items(), 
                               key=lambda x: x[1]['final_score'], 
                               reverse=True)
            
            # 상위 문서 반환
            result_docs = []
            for doc_idx, score_info in sorted_docs[:top_k]:
                doc = self.documents[doc_idx]
                print(f"📄 {doc.filename}: 점수 {score_info['max_score']:.3f} (청크 {score_info['count']}개)")
                print(f"   대표 청크: {score_info['best_chunk']}")
                result_docs.append(doc)
            
            print(f"✅ Semantic 검색 완료: {len(result_docs)}개 문서")
            return result_docs
            
        except Exception as e:
            print(f"⚠️ Semantic 검색 실패: {e}")
            return self._keyword_search(query, top_k)
    
    def _keyword_search(self, query: str, top_k: int = 3) -> List[Document]:
        """키워드 기반 검색 (백업)"""
        print(f"🔍 키워드 검색: '{query}' (총 {len(self.documents)}개 문서)")
        
        # 한글 키워드 추출
        query_lower = query.lower()
        query_keywords = self._extract_keywords(query_lower)
        print(f"🔑 추출된 키워드: {query_keywords}")
        
        # 각 문서에 대해 유사도 점수 계산
        scored_docs = []
        for doc in self.documents:
            score = self._calculate_similarity(query_keywords, doc.content.lower())
            print(f"📄 {doc.filename}: 점수 {score:.3f}")
            scored_docs.append((doc, score))
        
        # 점수 순으로 정렬하여 상위 k개 반환
        scored_docs.sort(key=lambda x: x[1], reverse=True)
        result = [doc for doc, score in scored_docs[:top_k]]
        print(f"✅ 키워드 검색 완료: {len(result)}개 문서")
        return result
    
    def _extract_keywords(self, text: str) -> List[str]:
        """키워드 추출"""
        # 한글, 영문, 숫자만 추출
        words = re.findall(r'[가-힣a-z0-9]+', text)
        # 2글자 이상 단어만 사용
        keywords = [word for word in words if len(word) >= 2]
        return keywords
    
    def _calculate_similarity(self, query_keywords: List[str], doc_content: str) -> float:
        """개선된 키워드 매칭 기반 유사도"""
        doc_keywords = self._extract_keywords(doc_content)
        
        if not query_keywords:
            return 0.1  # 기본 점수 부여
        
        if not doc_keywords:
            return 0.0
        
        # 매칭된 키워드 개수 기반 점수
        matches = 0
        partial_matches = 0
        
        for query_keyword in query_keywords:
            # 완전 매칭
            if query_keyword in doc_keywords:
                matches += 1
            # 부분 매칭 (키워드가 문서 내용에 포함)
            elif query_keyword in doc_content:
                partial_matches += 1
        
        # 기본 점수 계산
        if matches > 0:
            similarity = matches / len(query_keywords)
        elif partial_matches > 0:
            similarity = (partial_matches * 0.5) / len(query_keywords)
        else:
            # 일반적인 소나무재선충병 관련 문서면 기본 점수
            if any(word in doc_content for word in ['소나무', '재선충', '탐지', '병해', '산림']):
                similarity = 0.3
            else:
                similarity = 0.1
        
        # 특정 키워드에 높은 가중치 부여
        high_value_keywords = ['소나무재선충병', '증상', '초기', '탐지', '신뢰도', 'gps', '좌표', '방제', '피해목', '확산']
        for keyword in query_keywords:
            if keyword in high_value_keywords:
                if keyword in doc_content:
                    similarity += 0.4
        
        return min(similarity, 1.0)
    
    def _classify_question_type(self, keywords: List[str], question_text: str) -> str:
        """질문 유형을 분류"""
        
        # YOLO/모델 관련 질문
        yolo_patterns = ['yolo', '욜로', '모델', '탐지', '딥러닝', '학습', '훈련', '알고리즘', 'ai']
        if any(pattern in question_text for pattern in yolo_patterns):
            if any(kw in keywords for kw in ['어떻게', '방법', '하면', '좋을까', '개선']):
                return "yolo_detection"
        
        # 요약/정리 관련 질문  
        summary_patterns = ['정리', '요약', '간략', '등록', '참고', '문서', 'md', 'pdf']
        if any(pattern in question_text for pattern in summary_patterns):
            return "summary"
            
        # 증상 관련 질문
        symptom_patterns = ['증상', '초기', '특징', '변화', '어떤']
        if any(pattern in question_text for pattern in symptom_patterns):
            return "symptoms"
            
        # 방제 관련 질문
        control_patterns = ['방제', '처리', '대응', '조치', '관리']
        if any(pattern in question_text for pattern in control_patterns):
            return "control"
            
        # 기술적 질문
        technical_patterns = ['정확도', '성능', '개선', '최적화', '향상']
        if any(pattern in question_text for pattern in technical_patterns):
            return "technical"
            
        # 기본값
        return "general"
    
    def get_context_for_question(self, question: str) -> str:
        """질문에 대한 관련 컨텍스트 생성"""
        relevant_docs = self.simple_search(question, top_k=3)
        
        if not relevant_docs:
            return "관련 전문 지식을 찾을 수 없습니다."
        
        context_parts = []
        for i, doc in enumerate(relevant_docs, 1):
            # 문서 내용 요약 (처음 500자)
            content_preview = doc.content[:500].strip()
            if len(doc.content) > 500:
                content_preview += "..."
            
            context_part = f"""
📖 참고자료 {i} ({doc.metadata['source_file']}):
{content_preview}
"""
            context_parts.append(context_part)
        
        return "\n".join(context_parts)
    
    def generate_korean_response(self, question: str, context: str = "", data_analysis: str = "") -> str:
        """OpenAI GPT를 활용한 자연스러운 답변 생성"""
        
        print(f"💬 질문: {question}")
        
        # 1. 관련 정보 검색
        if not context:
            context = self.get_context_for_question(question)
        
        # 2. OpenAI로 답변 생성
        if self.openai_chain:
            try:
                print("🤖 OpenAI로 답변 생성 중...")
                print(f"📝 Context 길이: {len(context)} 글자")
                print(f"📝 Context 미리보기: {context[:200]}...")
                
                response = self.openai_chain.invoke({
                    "context": context,
                    "question": question
                })
                
                print(f"📤 OpenAI 원본 응답: '{response}'")
                print(f"📏 응답 길이: {len(response)} 글자")
                
                if response and len(response.strip()) > 10:
                    print(f"✅ OpenAI 답변 생성 완료 ({len(response)} 글자)")
                    return response.strip()
                else:
                    print(f"⚠️ OpenAI 응답이 너무 짧음: '{response}'")
                    
            except Exception as e:
                print(f"⚠️ OpenAI 답변 생성 실패: {e}")
                import traceback
                traceback.print_exc()
        
        # 3. Fallback: 지식베이스에서 직접 추출
        print("🔄 지식베이스에서 직접 답변 추출")
        return self._generate_dynamic_answer(question, context)
    
    def _generate_simple_gpt_response(self, question: str, context: str, data_analysis: str = "") -> str:
        """간단하고 확실한 GPT 응답 생성"""
        
        # 매우 간단한 프롬프트 구성
        if context and len(context.strip()) > 10:
            prompt = f"""질문: {question}

관련 정보:
{context[:500]}

위 정보를 참고해서 친근하고 자연스럽게 답변해주세요:"""
        else:
            prompt = f"""질문: {question}

소나무재선충병 전문가로서 친근하고 자연스럽게 답변해주세요:"""

        print(f"📝 프롬프트 길이: {len(prompt)} 글자")
        
        # 안전한 토큰화
        try:
            inputs = self.tokenizer.encode(
                prompt, 
                return_tensors='pt', 
                max_length=400,  # 더 짧게
                truncation=True
            )
            
            print(f"🤖 GPT 생성 중... (입력 토큰: {inputs.shape[1]})")
            
            # 매우 보수적인 생성 설정
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs,
                    max_new_tokens=100,  # 짧게
                    min_new_tokens=20,   # 최소 보장
                    temperature=0.8,
                    do_sample=True,
                    top_p=0.9,
                    pad_token_id=self.tokenizer.eos_token_id,
                    eos_token_id=self.tokenizer.eos_token_id,
                    repetition_penalty=1.1
                )
            
            # 응답 추출
            full_response = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 프롬프트 제거하고 응답만 추출
            if len(full_response) > len(prompt):
                response = full_response[len(prompt):].strip()
            else:
                response = ""
            
            print(f"📤 원본 응답 길이: {len(response)}")
            
            # 응답 정제
            if response:
                # 불완전한 문장 제거
                sentences = response.split('.')
                if len(sentences) > 1:
                    complete_sentences = [s.strip() for s in sentences[:-1] if len(s.strip()) > 10]
                    if complete_sentences:
                        response = '. '.join(complete_sentences) + '.'
                
                # 최소 길이 확인
                if len(response) >= 20:
                    print(f"✅ 정제된 응답: {len(response)} 글자")
                    return response
            
            print("⚠️ 응답이 너무 짧거나 비어있음")
            return ""
            
        except Exception as e:
            print(f"⚠️ GPT 생성 오류: {e}")
            return ""

    def _generate_natural_gpt_response(self, question: str, context: str, data_analysis: str = "") -> str:
        """자연스러운 GPT 응답 생성 - LangChain 스타일"""
        
        # 자연스러운 프롬프트 구성
        system_message = """당신은 친근하고 전문적인 소나무재선충병 AI 어시스턴트입니다. 
사용자와 자연스러운 대화하듯이 답변해주세요.

대화 스타일:
- 친근하고 도움이 되는 톤으로 답변
- 이모지 적절히 사용 (🌲🤖📊 등)
- 전문 용어는 쉽게 설명
- 구체적이고 실용적인 조언 제공
- 추가 질문을 유도하는 마무리"""

        # 컨텍스트가 있으면 포함
        knowledge_section = ""
        if context and "참고자료" in context:
            knowledge_section = f"\n\n### 참고 지식:\n{context}"

        # 데이터 분석이 있으면 포함  
        data_section = ""
        if data_analysis:
            data_section = f"\n\n### 현재 데이터:\n{data_analysis}"

        # 최종 프롬프트 구성
        full_prompt = f"""{system_message}

{knowledge_section}
{data_section}

### 사용자 질문:
{question}

### 전문가 답변:
"""

        print(f"📝 프롬프트 길이: {len(full_prompt)} 글자")
        
        # 토큰화 (attention_mask 포함)
        tokenized = self.tokenizer(
            full_prompt,
            return_tensors='pt',
            max_length=800,
            truncation=True,
            padding=True
        )

        input_ids = tokenized['input_ids']
        attention_mask = tokenized.get('attention_mask', None)

        # pad_token 설정(없는 경우 eos로 대체)
        if getattr(self.tokenizer, 'pad_token', None) is None:
            try:
                self.tokenizer.pad_token = self.tokenizer.eos_token
            except Exception:
                pass

        # 자연스러운 응답 생성
        with torch.no_grad():
            outputs = self.model.generate(
                input_ids,
                attention_mask=attention_mask,
                max_new_tokens=200,
                min_new_tokens=50,
                num_return_sequences=1,
                temperature=0.7,
                do_sample=True,
                top_p=0.92,
                repetition_penalty=1.1,
                pad_token_id=self.tokenizer.eos_token_id,
                eos_token_id=self.tokenizer.eos_token_id
            )
        
        # 응답 추출 및 정제
        full_response = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
        
        # 프롬프트 부분 제거하고 응답만 추출
        if "### 전문가 답변:" in full_response:
            response = full_response.split("### 전문가 답변:")[-1].strip()
        else:
            # 입력 프롬프트 길이만큼 제거
            response = full_response[len(full_prompt):].strip()
        
        # 응답 정제
        response = self._clean_gpt_response(response)
        
        print(f"🔍 GPT 원본 응답 길이: {len(response)} 글자")
        print(f"🔍 GPT 응답 미리보기: {response[:100]}...")
        
        # 최소 길이 확인 (기준을 낮춤)
        if len(response) < 15:
            print("⚠️ GPT 응답이 너무 짧음, 백업 사용")
            return self._generate_simple_fallback(question, context)
        
        # 의미있는 응답인지 확인
        if response.count(' ') < 5:  # 단어가 너무 적음
            print("⚠️ GPT 응답이 의미없음, 백업 사용")
            return self._generate_simple_fallback(question, context)
        
        return response
    
    def _clean_gpt_response(self, response: str) -> str:
        """GPT 응답 정제"""
        # 불완전한 문장 제거
        lines = response.split('\n')
        clean_lines = []
        
        for line in lines:
            line = line.strip()
            if len(line) > 5:  # 너무 짧은 줄 제거
                # 특수문자로만 이루어진 줄 제거
                if not line.replace('-', '').replace('*', '').replace('=', '').strip():
                    continue
                clean_lines.append(line)
        
        result = '\n'.join(clean_lines)
        
        # 마지막에 불완전한 문장 제거
        if result and not result[-1] in '.!?다요니까':
            sentences = result.split('.')
            if len(sentences) > 1:
                result = '.'.join(sentences[:-1]) + '.'
        
        return result.strip()
    
    def _generate_simple_fallback(self, question: str, context: str) -> str:
        """GPT 실패 시 자연스러운 대화형 응답"""
        
        if context and len(context.strip()) > 30:
            # 컨텍스트에서 핵심 정보만 추출
            clean_context = context.replace('\n', ' ').strip()
            
            # 너무 긴 경우 요약
            if len(clean_context) > 200:
                clean_context = clean_context[:200] + "..."
            
            # 질문 유형에 따른 자연스러운 응답
            if "확인" in question or "방법" in question:
                return f"재선충 확인방법에 대해 말씀드리면, {clean_context} 이런 정보가 있어요. 더 구체적으로 궁금한 부분이 있으면 알려주세요!"
            elif "탐지" in question:
                return f"탐지 관련해서는 {clean_context} 이런 내용이 있네요. 어떤 부분을 더 자세히 알고 싶으신가요?"
            else:
                return f"{question}에 대해서는 {clean_context} 이런 자료가 있어요. 추가로 궁금한 점이 있으시면 말씀해 주세요!"
        
        # 컨텍스트가 없을 때 질문 유형별 응답
        if "재선충" in question:
            return "소나무재선충병은 매개충을 통해 전파되는 질병이에요. 드론과 AI로 조기 탐지하는 것이 중요하죠. 구체적으로 어떤 부분이 궁금하신가요?"
        elif "탐지" in question or "YOLO" in question:
            return "YOLO 모델로 드론 영상에서 감염목을 찾아내고 있어요. 신뢰도에 따라 즉시 방제하거나 재조사를 하죠. 더 자세히 알고 싶은 부분이 있나요?"
        else:
            return f"'{question}' 관련해서 도움드릴 수 있어요! 소나무재선충병이나 AI 탐지에 대해 더 구체적으로 질문해주세요."

    def _generate_gpt_response(self, question: str, context: str, data_analysis: str = "") -> str:
        """간단하고 확실한 GPT 응답 생성"""
        
        try:
            # 아주 간단한 프롬프트 구성
            if context and len(context.strip()) > 20:
                clean_context = context[:300].replace('\n', ' ')
                prompt = f"질문: {question}\n참고정보: {clean_context}\n답변:"
            else:
                prompt = f"소나무재선충병 전문가로서 다음 질문에 답하세요.\n질문: {question}\n답변:"
            
            print(f"🎯 단순 프롬프트 (길이: {len(prompt)})")
            
            # 토큰화
            inputs = self.tokenizer(
                prompt,
                return_tensors='pt',
                max_length=600,
                truncation=True
            )
            
            print(f"🤖 GPT 생성 중... (입력토큰: {inputs['input_ids'].shape[1]})")
            
            # 단순한 생성 파라미터
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs['input_ids'],
                    max_new_tokens=150,
                    temperature=0.7,
                    do_sample=True,
                    pad_token_id=self.tokenizer.eos_token_id
                )
            
            # 응답 추출
            generated_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 프롬프트 부분 제거
            if "답변:" in generated_text:
                response = generated_text.split("답변:", 1)[-1].strip()
            else:
                response = generated_text[len(prompt):].strip()
            
            # 기본 정제
            response = response.split('\n')[0]  # 첫 번째 줄만
            response = response.strip()
            
            print(f"✅ GPT 응답: {response[:50]}... (길이: {len(response)})")
            
            if len(response) < 15:
                print("⚠️ 응답 너무 짧음, 백업 사용")
                return self._generate_simple_fallback(question, context)
                
            return response
            
        except Exception as e:
            print(f"❌ GPT 오류: {e}")
            return self._generate_simple_fallback(question, context)

    def _generate_smart_gpt_response(self, question: str, context: str, data_analysis: str = "") -> str:
        """GPT가 지식베이스를 활용해서 알아서 자연스럽게 답변"""
        
        # 완전 자연스러운 프롬프트 - 대화하듯이
        if context and len(context.strip()) > 20:
            prompt = f"""당신은 소나무재선충병 전문가입니다.

참고자료:
{context[:800]}

질문: {question}

위 자료를 바탕으로 친근하고 자연스럽게 답변해주세요:"""
        else:
            prompt = f"""당신은 소나무재선충병과 산림병해충 전문가입니다.

질문: {question}

전문 지식을 바탕으로 친근하고 자연스럽게 답변해주세요:"""

        try:
            inputs = self.tokenizer.encode_plus(
                prompt,
                return_tensors='pt',
                max_length=900,
                truncation=True,
                padding=True
            )
            
            print("🤖 자연스러운 GPT 응답 생성 중...")
            
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs['input_ids'],
                    attention_mask=inputs.get('attention_mask'),
                    max_new_tokens=200,
                    min_new_tokens=60,
                    temperature=0.85,
                    do_sample=True,
                    top_p=0.92,
                    repetition_penalty=1.1,
                    no_repeat_ngram_size=3,
                    pad_token_id=self.tokenizer.eos_token_id
                )
            
            full_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 프롬프트 부분 제거
            response = full_text.replace(prompt, "").strip()
            
            # 응답 정제
            if response.startswith(":"):
                response = response[1:].strip()
            
            # 너무 짧으면 자연스러운 대안
            if len(response) < 40:
                return self._generate_natural_fallback(question, context)
                
            return response
            
        except Exception as e:
            print(f"❌ GPT 오류: {e}")
            return self._generate_natural_fallback(question, context)

    def _generate_natural_fallback(self, question: str, context: str) -> str:
        """GPT 실패시에도 자연스럽게 답변"""
        
        if context and len(context.strip()) > 20:
            # 컨텍스트를 활용한 자연스러운 답변
            context_summary = context[:200]
            return f"{question}에 대해 말씀드리면,\n\n{context_summary}...\n\n이런 정보가 있네요! 더 자세한 설명이 필요하시면 언제든 말씀해주세요 😊"
        else:
            # 일반적인 도움 제공
            return f"{question}에 대해 궁금하시는군요! 소나무재선충병이나 AI 탐지 관련해서 더 구체적으로 질문해주시면 정확한 정보를 드릴 수 있어요. 어떤 부분이 가장 궁금하신가요?"
    
    def _build_natural_korean_prompt(self, question: str, context: str, data_analysis: str) -> str:
        """자연스러운 한국어 프롬프트 구성 - GPT 모델에 최적화"""
        
        prompt_parts = []
        
        # 간결하고 명확한 지시문
        prompt_parts.append("소나무재선충병 전문가로서 질문에 답하세요.")
        
        # 참고 자료가 있으면 요약해서 포함
        if context and len(context.strip()) > 20:
            # 너무 긴 컨텍스트는 요약
            clean_context = context[:400].replace('\n', ' ').strip()
            prompt_parts.append(f"참고자료: {clean_context}")
        
        # 명확한 질문과 응답 형식
        prompt_parts.append(f"질문: {question}")
        prompt_parts.append("답변:")
        
        return "\n".join(prompt_parts)
    
    def _clean_and_validate_response(self, response: str, question: str) -> str:
        """GPT 응답 정제 및 검증"""
        
        # 기본 정제
        response = response.strip()
        
        # 불필요한 접두사 제거
        prefixes_to_remove = [
            "소나무재선충병 전문가로서",
            "참고자료:",
            "질문:",
            "답변:",
            "사용자 질문:",
            "전문가 답변:",
        ]
        
        for prefix in prefixes_to_remove:
            if response.startswith(prefix):
                response = response[len(prefix):].strip()
        
        # 첫 번째 의미있는 문장 찾기
        lines = response.split('\n')
        clean_lines = []
        
        for line in lines:
            line = line.strip()
            if len(line) > 10 and not line.startswith(('참고자료', '질문:', '답변:')):
                clean_lines.append(line)
        
        if clean_lines:
            response = '\n'.join(clean_lines[:3])  # 최대 3줄
        
        # 중복 제거
        if question in response:
            response = response.replace(question, "").strip()
        
        return response
    
    def _generate_template_response(self, question: str, context: str, data_analysis: str) -> str:
        """지능형 템플릿 기반 한국어 응답 생성"""
        
        # 키워드 분석
        question_keywords = self._extract_keywords(question.lower())
        context_keywords = self._extract_keywords(context.lower()) if context else []
        
        response_parts = []
        
        # GPT 없이 템플릿 사용할 때는 간단하게
        print("📝 간단한 템플릿 기반 응답 생성 중...")
        
        # 컨텍스트 있으면 그것을 기반으로 자연스럽게 응답
        if context and len(context.strip()) > 10:
            response_parts.append("참고 자료를 바탕으로 답변드리겠습니다:")
            response_parts.append("")
            
            # 컨텍스트의 핵심 내용을 요약해서 제시
            context_summary = context[:400].replace("📖 참고자료", "•").strip()
            response_parts.append(context_summary)
            response_parts.append("")
            
            # 질문에 따른 추가 설명
            if any(kw in question_keywords for kw in ['yolo', '욜로', '모델', '탐지']):
                response_parts.append("YOLO 모델 기반 탐지의 핵심은 정확한 데이터셋 구축과 적절한 하이퍼파라미터 조정입니다.")
            elif any(kw in question_keywords for kw in ['증상', '특징']):
                response_parts.append("초기 증상을 정확히 파악하는 것이 효과적인 방제의 첫걸음입니다.")
            elif any(kw in question_keywords for kw in ['방제', '처리']):
                response_parts.append("신속한 초기 대응이 확산 방지의 핵심입니다.")
            
        else:
            response_parts.append("더 구체적인 정보를 제공해주시면 정확한 답변을 드릴 수 있습니다.")
            response_parts.append("예: 'YOLO 모델 학습 방법', '소나무재선충병 증상', '방제 처리 절차' 등")
        
        return "\n".join(response_parts) if response_parts else "죄송하지만 해당 질문에 대한 정보를 찾을 수 없습니다."

        
        # 데이터 요약
        if data_analysis:
            response_parts.append(f"📊 **현재 상황:** {data_analysis}")
            response_parts.append("")
        
        # 맥락 기반 상세 분석
        if context:
            if "신뢰도" in context and any(char.isdigit() for char in context):
                # 숫자 추출
                numbers = [int(s) for s in context.split() if s.isdigit()]
                if numbers:
                    confidence = max(numbers)
                    if confidence >= 80:
                        response_parts.append("✅ **평가:** 매우 우수한 탐지 결과입니다.")
                        response_parts.append("• 높은 신뢰도로 즉시 방제 조치 권장")
                    elif confidence >= 60:
                        response_parts.append("⚠️ **평가:** 양호한 탐지 결과입니다.")
                        response_parts.append("• 추가 확인 후 방제 조치 실시")
                    else:
                        response_parts.append("🔍 **평가:** 추가 검증이 필요합니다.")
                        response_parts.append("• 재탐지 또는 다른 방법 활용 권장")
                    response_parts.append("")
        
        # 질문 유형별 맞춤 답변
        if any(kw in question_keywords for kw in ['개선', '향상', '좋게', '높이']):
            response_parts.append("🎯 **정확도 개선 방안:**")
            response_parts.append("1. **데이터 품질 향상**")
            response_parts.append("   • 고해상도 이미지 사용")
            response_parts.append("   • 다양한 촬영 각도 확보")
            response_parts.append("2. **모델 재학습**")
            response_parts.append("   • 추가 학습 데이터 수집")
            response_parts.append("   • 하이퍼파라미터 최적화")
            response_parts.append("3. **앙상블 기법 활용**")
            response_parts.append("   • 여러 모델 결과 조합")
            response_parts.append("")
            response_parts.append("• 신뢰도 40-70%: 추가 현장 확인 후 방제")
            response_parts.append("• 신뢰도 40% 미만: 재탐지 또는 다른 방법 활용")
            response_parts.append("")
        
        if any(kw in question_keywords for kw in ['방제', '처리', '대응']):
            response_parts.append("🚨 **방제 처리 방안:**")
            response_parts.append("• 감염목 즉시 벌채 및 반출")
            response_parts.append("• 주변 500m 반경 예방 살포")
            response_parts.append("• GPS 좌표 기반 체계적 관리")
            response_parts.append("")
        
        if any(kw in question_keywords for kw in ['확산', '전파', '예방']):
            response_parts.append("🛡️ **확산 방지 대책:**")
            response_parts.append("• 매개충 활동 시기 집중 모니터링")
            response_parts.append("• 감염목 운반 경로 차단")
            response_parts.append("• 주변 지역 예방 조치 강화")
            response_parts.append("")
        
        if context and "참고자료" in context:
            response_parts.append("📚 **전문 지식 참고:**")
            response_parts.append("위 분석은 산림청 방제 가이드라인과 생태학적 특성을 기반으로 작성되었습니다.")
        
        return "\n".join(response_parts) if response_parts else "죄송하지만 해당 질문에 대한 구체적인 답변을 생성할 수 없습니다. 더 구체적인 질문을 해주시면 도움이 될 것 같습니다."
    
    def _postprocess_korean_response(self, response: str) -> str:
        """한국어 응답 후처리"""
        # 불완전한 문장 제거
        sentences = response.split('.')
        complete_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 10 and not sentence.endswith(('하고', '하여', '하면', '하는')):
                complete_sentences.append(sentence)
        
        if complete_sentences:
            result = '. '.join(complete_sentences)
            if not result.endswith('.'):
                result += '.'
            return result
        
        return response[:200] + "..." if len(response) > 200 else response

    def _generate_smart_gpt_response(self, question: str, context: str, data_analysis: str = "") -> str:
        """GPT가 스스로 판단해서 자연스러운 답변 생성"""
        
        try:
            # GPT가 스스로 판단할 수 있는 자연스러운 프롬프트
            if context and len(context.strip()) > 30:
                # 실제 지식 내용만 추출 (메타데이터 제거)
                clean_context = context.replace('📖 참고자료', '').replace('#', '').replace('*', '')
                clean_context = ' '.join([line.strip() for line in clean_context.split('\n') if len(line.strip()) > 10])[:500]
                
                prompt = f"""다음은 소나무재선충병에 대한 전문 자료입니다:
{clean_context}

질문: {question}

위 자료를 바탕으로 전문가답게 친근하게 답변해주세요:"""
            else:
                prompt = f"""당신은 소나무재선충병 전문가입니다. 
질문: {question}
전문적이면서도 이해하기 쉽게 답변해주세요:"""
            
            print(f"🎯 GPT 프롬프트 준비: {len(prompt)} 글자")
            
            # 간단한 토큰화
            inputs = self.tokenizer.encode(prompt, return_tensors='pt', max_length=900, truncation=True)
            
            print(f"🔥 GPT 생성 시작...")
            
            # 매우 안전한 생성 파라미터 (NaN/inf 방지)
            with torch.no_grad():
                try:
                    # 첫 번째 시도: 그리디 디코딩 (가장 안전)
                    outputs = self.model.generate(
                        inputs,
                        max_length=inputs.shape[1] + 100,
                        do_sample=False,  # 그리디 방식
                        pad_token_id=self.tokenizer.eos_token_id,
                        eos_token_id=self.tokenizer.eos_token_id,
                        num_return_sequences=1
                    )
                except:
                    # 두 번째 시도: 매우 보수적인 샘플링
                    outputs = self.model.generate(
                        inputs,
                        max_length=inputs.shape[1] + 80,
                        temperature=1.0,  # 기본 온도
                        do_sample=True,
                        top_k=50,  # top_p 대신 top_k 사용
                        pad_token_id=self.tokenizer.eos_token_id,
                        num_return_sequences=1
                    )
            
            # 전체 생성 텍스트
            generated_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 원래 프롬프트 부분 제거
            if len(generated_text) > len(prompt):
                answer = generated_text[len(prompt):].strip()
            else:
                answer = generated_text.strip()
            
            print(f"✨ GPT 생성 완료: '{answer[:50]}...'")
            
            # 의미있는 답변인지 확인
            if len(answer) > 20 and answer.count(' ') > 3:
                return answer[:300]  # 최대 300글자로 제한
            
        except Exception as e:
            print(f"💥 GPT 생성 오류: {e}")
            import traceback
            traceback.print_exc()
        
        return None
    
    def _clean_smart_response(self, response: str) -> str:
        """똑똑한 응답 정리"""
        # 원본 응답에서 불필요한 부분 제거
        response = response.strip()
        
        # 프롬프트 잔여물 제거
        unwanted_patterns = [
            '📖 참고자료',
            '관련 전문 자료:',
            '질문:',
            '답변:',
            '위 자료를 참고해서',
            '위 정보를 바탕으로',
            '#',
            '##'
        ]
        
        lines = response.split('\n')
        clean_lines = []
        
        for line in lines:
            line = line.strip()
            # 불필요한 패턴이 포함된 줄 제거
            skip_line = False
            for pattern in unwanted_patterns:
                if pattern in line:
                    skip_line = True
                    break
            
            if not skip_line and len(line) > 10:
                clean_lines.append(line)
                if len(clean_lines) >= 3:  # 최대 3줄로 제한
                    break
        
        if clean_lines:
            result = ' '.join(clean_lines)
            return result[:300] if len(result) > 300 else result
        else:
            # 모든 줄이 필터링되면 원본의 의미있는 부분만 추출
            words = response.split()
            meaningful_words = [w for w in words if len(w) > 2 and not w.startswith(('#', '📖'))]
            return ' '.join(meaningful_words[:30]) if meaningful_words else response[:100]
    
    def _generate_intelligent_answer(self, question: str, context: str) -> str:
        """지식베이스를 이용해 GPT처럼 똑똑하고 자연스러운 답변 생성"""
        
        if not context or len(context.strip()) < 20:
            return f"{question}에 대한 자료를 찾지 못했어요. 소나무재선충병, YOLO 탐지, 드론 촬영 등에 대해서 질문해보시면 도움드릴 수 있습니다!"
        
        # 컨텍스트에서 실제 유용한 정보 추출
        useful_content = self._extract_smart_content(context, question)
        
        if not useful_content:
            return f"{question}에 대한 관련 자료는 있지만, 좀 더 구체적으로 질문해주시면 더 정확한 답변을 드릴 수 있을 것 같아요!"
        
        # 질문 유형에 따라 자연스러운 답변 구성
        return self._compose_natural_answer(question, useful_content)
    
    def _extract_smart_content(self, context: str, question: str) -> str:
        """컨텍스트에서 실제 유용한 내용만 간단하게 추출"""
        
        # 메타데이터 제거 패턴
        clean_text = context
        for pattern in ['📖', '참고자료', '대표 청크:', '점수', 'filepath:', '#', '*']:
            clean_text = clean_text.replace(pattern, '')
        
        # 줄 단위로 나누고 의미있는 문장들만 추출
        lines = [line.strip() for line in clean_text.split('\n') if line.strip()]
        useful_lines = []
        
        for line in lines:
            # 너무 짧거나 메타데이터인 줄 제외
            if len(line) < 15 or line.startswith('http') or '.pdf' in line or '.jpg' in line:
                continue
            
            # 질문과 관련된 키워드가 있는 줄 우선
            if question and any(keyword in line.lower() for keyword in question.lower().split()):
                useful_lines.insert(0, line)
            else:
                useful_lines.append(line)
            
            # 최대 3개 문장만
            if len(useful_lines) >= 3:
                break
        
        if useful_lines:
            result = '. '.join(useful_lines[:2])  # 최대 2개 문장
            return result[:300] + ('...' if len(result) > 300 else '')
        
        return context[:200] + '...' if len(context) > 200 else context
    
    def _compose_natural_answer(self, question: str, content: str) -> str:
        """질문과 내용을 바탕으로 자연스럽고 도움이 되는 답변 구성"""
        
        # [하드코딩 완전 제거 - LLM이 자유롭게 답변]
        # 이 함수는 더 이상 사용되지 않음 - LangChain이 자동 처리
        return content[:200] if content else "정보를 찾을 수 없습니다."

    def _extract_real_content(self, context: str) -> str:
        """컨텍스트에서 실제 유용한 내용만 추출 (메타데이터 제거)"""
        
        # 불필요한 메타데이터 패턴들
        skip_patterns = [
            '📖 참고자료', '# 무인항공기를', '* 한국교통연구원', 
            '** 종신회원', '.pdf', '.jpg', '.png', 'Fig.', 'Table.'
        ]
        
        # 줄 단위로 분리해서 실제 내용만 추출
        lines = context.split('\n')
        content_lines = []
        
        for line in lines:
            line = line.strip()
            
            # 메타데이터 줄 건너뛰기
            if any(skip in line for skip in skip_patterns):
                continue
                
            # 의미있는 내용이 있는 줄만 선택
            if len(line) > 15:
                content_lines.append(line)
                if len(content_lines) >= 4:  # 최대 4줄
                    break
        
        if content_lines:
            return '\n'.join(content_lines)
        else:
            # 줄 기준으로 안 되면 문장 기준으로 시도
            sentences = context.replace('\n', ' ').split('.')
            for sentence in sentences[:3]:  # 처음 3개 문장만
                sentence = sentence.strip()
                if len(sentence) > 20 and not any(skip in sentence for skip in skip_patterns):
                    return sentence + '.'
            
            return "관련 전문 자료의 내용을 확인했습니다."

    def _generate_gpt_knowledge_response(self, question: str, context: str) -> str:
        """GPT 모델 자체의 지식을 활용한 답변 생성"""
        
        try:
            # GPT가 자체 지식으로 답변할 수 있도록 하는 프롬프트
            if context and len(context.strip()) > 50:
                # 지식베이스 정보가 있으면 보완해서 답변
                prompt = f"""질문: {question}

참고 정보가 있지만 부족합니다. 당신의 지식을 활용해서 이 질문에 대해 전문적이고 친근하게 답변해주세요.

답변:"""
            else:
                # 지식베이스에 없으면 GPT 자체 지식으로 답변
                prompt = f"""질문: {question}

위 질문에 대해 당신이 알고 있는 지식을 바탕으로 친근하고 전문적으로 답변해주세요.

답변:"""
            
            print(f"🎯 GPT 자체 지식 활용 중...")
            
            # 간단한 토큰화
            inputs = self.tokenizer.encode(prompt, return_tensors='pt', max_length=700, truncation=True)
            
            # 안전한 생성
            with torch.no_grad():
                try:
                    outputs = self.model.generate(
                        inputs,
                        max_length=inputs.shape[1] + 100,
                        temperature=0.7,
                        do_sample=True,
                        top_k=50,
                        pad_token_id=self.tokenizer.eos_token_id
                    )
                except:
                    # 더 안전한 그리디 방식
                    outputs = self.model.generate(
                        inputs,
                        max_length=inputs.shape[1] + 80,
                        do_sample=False,
                        pad_token_id=self.tokenizer.eos_token_id
                    )
            
            # 응답 추출
            generated_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 프롬프트 제거
            if len(generated_text) > len(prompt):
                answer = generated_text[len(prompt):].strip()
            else:
                return None
            
            # 의미있는 답변인지 확인 (반복문자 제거)
            if len(answer) > 30 and not ('!!!' in answer or answer.count(answer[0]) > 10):
                # 답변 정리 (최대 300글자)
                if len(answer) > 300:
                    answer = answer[:300] + "..."
                return answer
                
        except Exception as e:
            print(f"💥 GPT 지식 활용 오류: {e}")
        
        return None

    def _generate_pure_llm_response(self, question: str, context: str) -> str:
        """순수하게 LLM 자체 지식만으로 답변 생성 (하드코딩 없음)"""
        
        if not self.model or not self.tokenizer:
            return f"{question}에 대한 답변을 위해 모델을 로딩 중입니다. 잠시 후 다시 시도해주세요."
        
        try:
            # 순수 LLM 답변을 위한 자연스러운 프롬프트
            if context and len(context.strip()) > 50:
                prompt = f"""사용자 질문: {question}

관련 참고 정보:
{context[:400]}

위 정보를 참고하여 질문에 대해 친근하고 전문적으로 답변해주세요:"""
            else:
                prompt = f"""사용자 질문: {question}

위 질문에 대해 친근하고 전문적으로 답변해주세요:"""
            
            print(f"🎯 순수 LLM 프롬프트 준비 완료")
            
            # 토큰화
            inputs = self.tokenizer.encode(prompt, return_tensors='pt', max_length=800, truncation=True)
            
            # 안전하고 자연스러운 생성
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs,
                    max_length=inputs.shape[1] + 200,
                    min_length=inputs.shape[1] + 50,
                    temperature=0.8,
                    do_sample=True,
                    top_p=0.9,
                    repetition_penalty=1.1,
                    no_repeat_ngram_size=3,
                    pad_token_id=self.tokenizer.eos_token_id,
                    eos_token_id=self.tokenizer.eos_token_id
                )
            
            # 응답 추출
            generated_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 프롬프트 제거하고 순수 답변만 추출
            if len(generated_text) > len(prompt):
                answer = generated_text[len(prompt):].strip()
            else:
                return self._generate_fallback_response(question, context)
            
            # 답변 품질 검증
            if len(answer) > 30 and self._is_valid_response(answer):
                return answer[:500]  # 최대 500글자
            else:
                return self._generate_fallback_response(question, context)
                
        except Exception as e:
            print(f"💥 LLM 생성 오류: {e}")
            return self._generate_fallback_response(question, context)
    
    def _is_valid_response(self, response: str) -> bool:
        """생성된 응답이 유효한지 검증"""
        # 반복 문자나 무의미한 토큰 체크
        if response.count('!') > 5 or response.count('?') > 3:
            return False
        if any(char * 10 in response for char in 'abcdefghijklmnopqrstuvwxyz'):
            return False
        return True
    
    def _generate_fallback_response(self, question: str, context: str) -> str:
        """LLM 실패시 최소한의 도움이 되는 응답"""
        if context and len(context.strip()) > 30:
            clean_context = self._extract_smart_content(context, question)
            if clean_context:
                return f"{question}에 대한 관련 정보를 찾았습니다:\n\n{clean_context}\n\n더 구체적인 질문이 있으시면 말씀해 주세요!"
        
        return f"{question}에 대한 답변을 생성하는 중 문제가 발생했습니다. 다른 방식으로 질문해 보시거나, 좀 더 구체적으로 말씀해 주시겠어요?"

    def _generate_pure_gpt_response(self, question: str, context: str) -> str:
        """순수 GPT 모델 지식만 활용한 자연스러운 답변 생성 (하드코딩 없음)"""
        
        try:
            # 자연스러운 한국어 대화 프롬프트 (컨텍스트 활용)
            if context and len(context.strip()) > 30:
                # 지식베이스 정보가 있으면 참고해서 답변
                clean_context = context.replace('📖', '').replace('#', '').strip()[:300]
                prompt = f"""사용자 질문: {question}
                
참고 자료: {clean_context}

위 질문에 대해 참고 자료와 당신의 지식을 종합해서 친근하고 전문적으로 답변해주세요:

답변:"""
            else:
                # 지식베이스 정보가 없으면 GPT 자체 지식으로만 답변
                prompt = f"""사용자 질문: {question}

위 질문에 대해 당신이 알고 있는 지식을 바탕으로 친근하고 도움이 되는 답변을 해주세요:

답변:"""
            
            print(f"🎯 순수 GPT 답변 생성 시작...")
            
            # 토큰화
            inputs = self.tokenizer.encode(prompt, return_tensors='pt', max_length=900, truncation=True)
            
            # 더 나은 생성 파라미터 (한국어 친화적)
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs,
                    max_length=inputs.shape[1] + 200,  # 더 긴 답변 허용
                    temperature=0.8,  # 창의적 답변
                    do_sample=True,
                    top_p=0.9,
                    top_k=40,
                    repetition_penalty=1.1,
                    pad_token_id=self.tokenizer.eos_token_id,
                    eos_token_id=self.tokenizer.eos_token_id
                )
            
            # 전체 생성 텍스트 추출
            full_response = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 프롬프트 부분 제거
            if len(full_response) > len(prompt):
                answer = full_response[len(prompt):].strip()
            else:
                return None
            
            # 응답 품질 확인
            if len(answer) > 30 and self._is_meaningful_response(answer):
                return answer[:500]  # 최대 500자로 제한
                
        except Exception as e:
            print(f"💥 순수 GPT 답변 생성 오류: {e}")
        
        return None
    
    def _generate_langchain_response(self, question: str, context: str) -> str:
        """LangChain을 활용한 한국어 자연 답변 생성"""
        
        # 프롬프트 템플릿 설정 (한국어 대화형)
        if context and len(context.strip()) > 30:
            # 지식베이스 정보 활용
            prompt_template = PromptTemplate(
                input_variables=["question", "context"],
                template="""당신은 친근하고 전문적인 AI 어시스턴트입니다.

참고 정보:
{context}

사용자 질문: {question}

위 질문에 대해 참고 정보와 당신의 지식을 바탕으로 도움이 되는 답변을 해주세요. 자연스럽고 친근한 어조로 답변해주세요:

답변:"""
            )
        else:
            # GPT 자체 지식만 활용
            prompt_template = PromptTemplate(
                input_variables=["question"],
                template="""당신은 친근하고 전문적인 AI 어시스턴트입니다.

사용자 질문: {question}

위 질문에 대해 당신이 알고 있는 지식을 바탕으로 도움이 되는 답변을 해주세요. 자연스럽고 친근한 어조로 답변해주세요:

답변:"""
            )
        
        try:
            # LangChain 새로운 방식으로 체인 생성
            chain = prompt_template | self.langchain_llm
            
            # 답변 생성
            if context and len(context.strip()) > 30:
                clean_context = context.replace('📖', '').replace('#', '').strip()[:400]
                response = chain.invoke({"question": question, "context": clean_context})
            else:
                response = chain.invoke({"question": question})
            
            # 응답 정리 (불필요한 부분 제거)
            if response:
                response = str(response).strip()
                # 프롬프트 반복 제거
                if "답변:" in response:
                    response = response.split("답변:")[-1].strip()
                
                # 의미있는 응답인지 확인
                if len(response) > 30 and self._is_meaningful_response(response):
                    return response[:600]  # 최대 600자로 제한
                    
        except Exception as e:
            print(f"💥 LangChain 답변 생성 오류: {e}")
        
        return None
    
    def _create_natural_gpt_response(self, question: str, context: str) -> str:
        """GPT의 자연스러운 대화능력과 전문지식을 활용한 응답 생성"""
        
        # 질문에 따라 GPT가 자연스럽게 대답할 수 있도록 대화형 프롬프트 구성
        if context and len(context.strip()) > 30:
            conversation_prompt = f"""안녕하세요! 저는 AI와 산림병해 전문가입니다. 

"{question}"

이 질문에 대해서 제가 알고 있는 지식을 바탕으로 설명드릴게요. 
참고로 관련 자료에서는 이런 내용이 있네요: {context[:200]}...

자세히 설명드리면:"""
        else:
            conversation_prompt = f"""안녕하세요! 저는 AI와 산림병해 전문가입니다.

"{question}"

이 질문에 대해서 제가 알고 있는 전문 지식을 바탕으로 친근하게 설명드릴게요:"""
        
        # 더 자유로운 생성을 위한 시도
        try:
            inputs = self.tokenizer.encode(conversation_prompt, return_tensors='pt', max_length=400, truncation=True)
            attention_mask = (inputs != self.tokenizer.pad_token_id).long()
            
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs,
                    attention_mask=attention_mask,
                    max_length=inputs.shape[1] + 300,
                    temperature=1.0,  # 더 창의적
                    do_sample=True,
                    top_p=0.95,
                    top_k=50,
                    repetition_penalty=1.3,
                    pad_token_id=self.tokenizer.eos_token_id,
                    no_repeat_ngram_size=3
                )
            
            full_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 프롬프트 제거하고 GPT 답변 부분만 추출
            if "설명드리면:" in full_text:
                answer = full_text.split("설명드리면:")[-1].strip()
            elif "설명드릴게요:" in full_text:
                answer = full_text.split("설명드릴게요:")[-1].strip()
            else:
                answer = full_text[len(conversation_prompt):].strip()
            
            # 답변 품질 확인
            if len(answer) > 50 and any(ending in answer for ending in ['니다', '어요', '습니다', '에요']):
                return f"🤖 **전문가 답변:**\n\n{answer}"
            
        except Exception as e:
            logger.warning(f"자연스러운 GPT 응답 생성 실패: {e}")
        
        return None
    
    def _langchain_free_response(self, question: str, context: str = "") -> str:
        """지식베이스 기반 자연스러운 답변 생성"""
        # API 키 없이 지식베이스만으로 답변
        return self._extract_natural_answer(question, context)
    
    def _generate_free_llm_response(self, question: str, context: str = "") -> str:
        """LLM이 완전히 자유롭게 답변 생성 - 하드코딩 절대 금지"""
        
        # 매우 심플한 대화형 프롬프트 - LLM이 자기 지식으로 자유롭게 답변
        if context and len(context.strip()) > 20:
            # 지식베이스 정보가 있으면 참고만
            prompt = f"""### 사용자 질문:
{question}

### 참고 자료 (선택적):
{context[:400]}

### AI 답변:
"""
        else:
            # 지식베이스 없으면 LLM 자체 지식만
            prompt = f"""### 사용자 질문:
{question}

### AI 답변:
"""
        
        try:
            print("🤖 LLM 자유 답변 생성 중...")
            
            # 토큰화 (CPU에서 실행)
            inputs = self.tokenizer(
                prompt,
                return_tensors='pt',
                max_length=600,
                truncation=True,
                padding=True
            ).to('cpu')  # CPU 명시
            
            # 자유로운 생성 설정
            with torch.no_grad():
                outputs = self.model.generate(
                    input_ids=inputs['input_ids'].to('cpu'),
                    attention_mask=inputs.get('attention_mask').to('cpu') if inputs.get('attention_mask') is not None else None,
                    max_new_tokens=250,
                    min_new_tokens=30,
                    temperature=0.7,
                    do_sample=True,
                    top_p=0.92,
                    top_k=50,
                    repetition_penalty=1.2,
                    no_repeat_ngram_size=3,
                    pad_token_id=self.tokenizer.pad_token_id or self.tokenizer.eos_token_id,
                    eos_token_id=self.tokenizer.eos_token_id
                )
            
            # 답변 추출
            generated = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 프롬프트 제거
            if "### AI 답변:" in generated:
                answer = generated.split("### AI 답변:")[-1].strip()
            elif len(generated) > len(prompt):
                answer = generated[len(prompt):].strip()
            else:
                answer = generated.strip()
            
            # 기본 정제
            if answer and len(answer) > 20:
                # 너무 긴 답변은 적절히 자르기
                if len(answer) > 500:
                    sentences = answer.split('.')
                    answer = '. '.join(sentences[:5]) + '.'
                
                print(f"✅ LLM 답변 성공 ({len(answer)}자)")
                return answer
            
        except Exception as e:
            print(f"❌ LLM 생성 오류: {e}")
        
        # 실패시 지식베이스만 반환
        return context[:300] if context else "죄송합니다. 답변을 생성할 수 없습니다."
    
    def _generate_smart_answer_with_gpt_knowledge(self, question: str, context: str) -> str:
        """[DEPRECATED] GPT의 자체 지식을 최대한 활용하여 전문적인 답변 생성 - 하드코딩 없이"""
        
        # GPT가 가진 전문 지식을 자유롭게 활용하도록 하는 간단한 프롬프트
        prompt = f"""당신은 AI, 컴퓨터 비전, 산림병해 전문가입니다. 

질문: {question}

위 질문에 대해 당신이 알고 있는 전문 지식을 바탕으로 자세하고 실용적인 답변을 한국어로 해주세요. 
구체적인 방법, 단계, 원리 등을 포함해서 설명해주세요.

답변:"""
        
        try:
            # attention mask와 함께 토큰화
            encoding = self.tokenizer(
                prompt, 
                return_tensors='pt', 
                max_length=700, 
                truncation=True,
                padding=True
            )
            
            with torch.no_grad():
                outputs = self.model.generate(
                    input_ids=encoding['input_ids'],
                    attention_mask=encoding['attention_mask'],
                    max_length=encoding['input_ids'].shape[1] + 200,
                    temperature=0.8,
                    do_sample=True,
                    top_p=0.9,
                    top_k=50,
                    repetition_penalty=1.1,
                    pad_token_id=self.tokenizer.eos_token_id,
                    eos_token_id=self.tokenizer.eos_token_id
                )
            
            # 전체 생성 텍스트 추출
            full_response = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 프롬프트 부분 제거하고 답변만 추출
            if "답변:" in full_response:
                answer = full_response.split("답변:")[-1].strip()
            elif len(full_response) > len(prompt):
                answer = full_response[len(prompt):].strip()
            else:
                return None
            
            # 응답 품질 확인 및 정리
            if len(answer) > 50 and self._is_meaningful_response(answer):
                # 불완전한 문장 제거
                sentences = answer.split('.')
                complete_sentences = []
                for sentence in sentences:
                    sentence = sentence.strip()
                    if len(sentence) > 20 and not sentence.endswith('...'):
                        complete_sentences.append(sentence)
                    if len(complete_sentences) >= 4:  # 최대 4문장
                        break
                
                if complete_sentences:
                    return '. '.join(complete_sentences) + '.'
                else:
                    return answer[:400]  # 최대 400자
                    
        except Exception as e:
            print(f"💥 GPT 전문 답변 생성 오류: {e}")
        
        return None
    
    def _is_meaningful_response(self, response: str) -> bool:
        """의미있는 응답인지 확인 (반복 문자, 무의미한 토큰 제거)"""
        # 같은 문자나 토큰의 과도한 반복 확인
        if len(set(response.replace(' ', ''))) < 5:  # 너무 적은 종류의 문자
            return False
        
        # 무의미한 반복 패턴 확인
        if '!!!' in response or '???' in response:
            return False
            
        # 특정 문자가 과도하게 반복되는지 확인
        for char in response:
            if response.count(char) > len(response) * 0.3:  # 30% 이상 같은 문자
                return False
        
        return True
    
    def _generate_natural_korean_response(self, question: str, context: str) -> str:
        """한국어 특화 모델을 사용한 자연스러운 답변 생성 (하드코딩 없음)"""
        
        # 한국어 질문 프롬프트 구성 - 모델이 자유롭게 답변할 수 있도록
        if context and len(context.strip()) > 30:
            prompt = f"""사용자: {question}

참고 정보: {context[:300]}

AI: 네, 질문에 대해 답변드리겠습니다. """
        else:
            prompt = f"""사용자: {question}

AI: 네, 그 질문에 대해 설명드리겠습니다. """
        
        # 여러 한국어 생성 방법 시도
        methods = [
            self._try_korean_specialized_model,
            self._try_langchain_korean_generation,
            self._try_improved_korean_generation
        ]
        
        for method in methods:
            try:
                result = method(prompt, question, context)
                if result and len(result.strip()) > 30 and self._is_valid_korean_response(result):
                    return result
            except Exception as e:
                print(f"⚠️ 한국어 생성 방법 실패: {e}")
                continue
        
        return None
    
    def _try_korean_specialized_model(self, prompt: str, question: str, context: str) -> str:
        """한국어 특화 모델로 답변 생성"""
        if not self.model or not self.tokenizer:
            return None
            
        try:
            # 한국어 모델에 최적화된 생성 설정
            encoding = self.tokenizer(
                prompt,
                return_tensors='pt',
                max_length=800,
                truncation=True,
                padding=True
            )
            
            with torch.no_grad():
                outputs = self.model.generate(
                    input_ids=encoding['input_ids'],
                    attention_mask=encoding['attention_mask'],
                    max_length=encoding['input_ids'].shape[1] + 300,
                    temperature=0.7,
                    do_sample=True,
                    top_p=0.9,
                    top_k=50,
                    repetition_penalty=1.2,
                    no_repeat_ngram_size=3,
                    pad_token_id=self.tokenizer.eos_token_id,
                    eos_token_id=self.tokenizer.eos_token_id
                )
            
            generated_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # AI 답변 부분만 추출
            if "AI: " in generated_text:
                answer = generated_text.split("AI: ")[-1].strip()
                if answer.startswith("네, "):
                    answer = answer[3:]  # "네, " 제거
                return answer
                
        except Exception as e:
            print(f"한국어 특화 모델 생성 실패: {e}")
        
        return None
    
    def _try_langchain_korean_generation(self, prompt: str, question: str, context: str) -> str:
        """LangChain을 통한 한국어 답변 생성"""
        if not self.langchain_llm:
            return None
            
        try:
            # 간단한 한국어 프롬프트
            simple_prompt = f"{question}에 대해 친절하고 전문적으로 답변해 주세요."
            
            response = self.langchain_llm.invoke(simple_prompt)
            if response and len(response.strip()) > 20:
                return response.strip()
                
        except Exception as e:
            print(f"LangChain 한국어 생성 실패: {e}")
        
        return None
    
    def _try_improved_korean_generation(self, prompt: str, question: str, context: str) -> str:
        """개선된 한국어 답변 생성"""
        if not self.model or not self.tokenizer:
            return None
            
        try:
            # 더 자연스러운 대화형 프롬프트
            conversation_prompt = f"전문가와의 대화:\n\n질문: {question}\n전문가: "
            
            inputs = self.tokenizer.encode(conversation_prompt, return_tensors='pt', max_length=600, truncation=True)
            attention_mask = (inputs != self.tokenizer.pad_token_id).long()
            
            with torch.no_grad():
                outputs = self.model.generate(
                    inputs,
                    attention_mask=attention_mask,
                    max_length=inputs.shape[1] + 250,
                    temperature=0.8,
                    do_sample=True,
                    top_p=0.95,
                    repetition_penalty=1.1,
                    pad_token_id=self.tokenizer.eos_token_id
                )
            
            full_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 전문가 답변 부분만 추출
            if "전문가: " in full_text:
                answer = full_text.split("전문가: ")[-1].strip()
                return answer
                
        except Exception as e:
            print(f"개선된 한국어 생성 실패: {e}")
        
        return None
    
    def _is_valid_korean_response(self, response: str) -> bool:
        """한국어 답변이 유효한지 확인"""
        if not response or len(response.strip()) < 10:
            return False
            
        # 한국어 문장 종결어미 확인
        korean_endings = ['다', '요', '니다', '습니다', '어요', '아요', '에요', '예요']
        has_korean_ending = any(response.strip().endswith(ending) for ending in korean_endings)
        
        # 한글 비율 확인
        korean_chars = sum(1 for char in response if '\uAC00' <= char <= '\uD7A3')
        korean_ratio = korean_chars / len(response) if len(response) > 0 else 0
        
        return has_korean_ending and korean_ratio > 0.3
    
    def _generate_dynamic_answer(self, question: str, context: str) -> str:
        """지식베이스 정보를 바탕으로 동적으로 답변 생성 (템플릿 없음)"""
        
        if not context or len(context.strip()) < 20:
            return f"{question}에 대한 정보를 찾지 못했습니다."
        
        # 강력한 메타데이터 필터
        skip_words = [
            'Received', 'revised', 'accepted', 'Copyright', 'Creative Commons', 
            'Attribution', 'Open Access', 'filepath:', 'koti.re.kr', '연구원', 
            '교수', '공학석사', '공학박사', '📖', '참고자료', '점수', '대표 청크',
            'Korea Transport', 'Fig.', 'Table', 'Graduate School', '.pdf', '.md'
        ]
        
        # 핵심 정보만 추출
        clean_lines = []
        for line in context.split('\n'):
            line = line.strip()
            
            # 빈 줄
            if not line:
                continue
            
            # 메타데이터 포함 줄 스킵
            if any(skip in line for skip in skip_words):
                continue
            
            # 제목 형식 줄 스킵  
            if line.startswith('#') or line.startswith('*') or line.startswith('('):
                continue
            
            # 너무 짧은 줄만 스킵
            if len(line) < 20:
                continue
            
            # 의미있는 내용
            clean_lines.append(line)
            if len(clean_lines) >= 4:
                break
        
        if clean_lines:
            return '\n\n'.join(clean_lines)
        
        return "관련 정보를 찾았으나 적절한 답변을 추출하지 못했습니다."
    
    def _pure_llm_knowledge_response(self, question: str) -> str:
        """[DEPRECATED] LLM 자체 지식으로만 답변 - 지식베이스/외부자료 절대 사용 안함"""
        
        # 심플한 프롬프트
        prompt = f"""질문: {question}

위 질문에 대해 한국어로 자연스럽게 답변해주세요:"""
        
        try:
            inputs = self.tokenizer(
                prompt,
                return_tensors='pt',
                max_length=150,
                truncation=True
            ).to('cpu')
            
            with torch.no_grad():
                outputs = self.model.generate(
                    input_ids=inputs['input_ids'],
                    attention_mask=inputs.get('attention_mask'),
                    max_new_tokens=250,
                    min_new_tokens=40,
                    temperature=0.8,
                    do_sample=True,
                    top_p=0.95,
                    repetition_penalty=1.3,
                    pad_token_id=self.tokenizer.eos_token_id,
                    eos_token_id=self.tokenizer.eos_token_id
                )
            
            response = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # 프롬프트 제거
            if len(response) > len(prompt):
                answer = response[len(prompt):].strip()
            else:
                answer = response
            
            if answer and len(answer) > 20:
                return answer
                
        except Exception as e:
            print(f"❌ LLM 생성 오류: {e}")
        
        return "죄송합니다. 답변을 생성할 수 없습니다."
    
    def _extract_natural_answer(self, question: str, context: str) -> str:
        """[DEPRECATED] 지식베이스에서 자연스러운 답변 추출 (하드코딩 없이 동적 생성)"""
        
        if not context or len(context.strip()) < 20:
            return f"{question}에 대한 정보를 찾을 수 없습니다. 다른 질문을 해주세요."
        
        # 강력한 메타데이터 제거
        skip_patterns = [
            'Received', 'revised', 'accepted', 'Creative Commons', 'Attribution',
            'Copyright', '한국교통연구원', '종신회원', '교신저자', '공학석사', '공학박사',
            'filepath:', '.pdf', '.md', '.txt', '.jpg', '.png',
            '📖', '참고자료', '점수', 'This is an Open Access',
            'The Korea Transport', 'koti.re.kr', 'Fig.', 'Table.',
            '무인항공기를이용한딥러닝', '연구원', '대학교', '교수'
        ]
        
        lines = context.split('\n')
        key_info = []
        
        for line in lines:
            line = line.strip()
            
            # 빈 줄 건너뛰기
            if not line:
                continue
            
            # 불필요한 패턴 건너뛰기
            if any(pattern in line for pattern in skip_patterns):
                continue
            
            # 메타데이터로 보이는 줄 건너뛰기
            if line.startswith('#') or line.startswith('*') or line.startswith('('):
                continue
            
            # 숫자만 있거나 매우 짧은 줄 건너뛰기
            if line.replace(' ', '').replace('.', '').replace('-', '').replace(':', '').isdigit():
                continue
            
            if len(line) < 15:
                continue
            
            # 의미 있는 실제 내용만 선택
            if any(keyword in line for keyword in ['탐지', '분석', '학습', '데이터', '모델', '방법', '확산', '감염', '소나무', 'YOLO', '드론', '신뢰도', 'GPS']):
                key_info.append(line)
                if len(key_info) >= 4:  # 최대 4줄
                    break
        
        if key_info:
            # 질문 유형에 따라 적절한 인트로 추가
            intro = self._generate_intro(question)
            info_text = '\n\n'.join(key_info)
            return f"{intro}\n\n{info_text}\n\n더 궁금한 점이 있으시면 말씀해주세요!"
        
        # Fallback - 컨텍스트 전체에서 핵심 키워드가 있는 부분 추출
        sentences = context.split('.')
        relevant_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 20 and not any(pattern in sentence for pattern in skip_patterns):
                if any(keyword in sentence for keyword in ['탐지', '분석', '학습', '데이터', '모델', '방법', '확산', '소나무', 'YOLO', '드론']):
                    relevant_sentences.append(sentence)
                    if len(relevant_sentences) >= 3:
                        break
        
        if relevant_sentences:
            intro = self._generate_intro(question)
            return f"{intro}\n\n{'. '.join(relevant_sentences)}.\n\n더 궁금한 점이 있으시면 말씀해주세요!"
        
        return f"{question}에 대한 구체적인 정보를 찾기 어렵습니다. 다른 방식으로 질문해주시겠어요?"
    
    def _generate_intro(self, question: str) -> str:
        """질문에 맞는 자연스러운 인트로 생성 (하드코딩 없이 동적)"""
        q_lower = question.lower()
        
        if '뭐' in q_lower or '무엇' in q_lower or '이란' in q_lower:
            return f"**{question}**\n\n알려드리겠습니다:"
        elif '어떻게' in q_lower or '방법' in q_lower:
            return f"**{question}**\n\n다음과 같은 방법이 있습니다:"
        elif '왜' in q_lower or '이유' in q_lower:
            return f"**{question}**\n\n이유를 설명드리겠습니다:"
        else:
            return f"**{question}**\n\n관련 정보입니다:"
    
    def _generate_knowledge_only_answer(self, question: str, context: str) -> str:
        """[DEPRECATED] 지식베이스만 사용 - _extract_natural_answer로 대체됨"""
        return self._extract_natural_answer(question, context)

# 전역 RAG 인스턴스
_rag_instance: Optional[SimpleRAG] = None

def get_rag_system() -> SimpleRAG:
    """RAG 시스템 싱글톤 인스턴스 반환"""
    global _rag_instance
    
    if _rag_instance is None:
        # 프로젝트 루트 디렉토리에서 knowledge_base 찾기
        current_dir = Path(__file__).parent.parent  # api 폴더에서 프로젝트 루트로
        knowledge_base_path = current_dir / "knowledge_base"
        
        _rag_instance = SimpleRAG(str(knowledge_base_path))
    
    return _rag_instance